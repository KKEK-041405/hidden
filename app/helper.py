import PyPDF2
import docx

def convert():
    # Open the PDF file
    pdf_file = open('C:\\Users\\Sandeep\\Documents\\kkek\\sk\\documents\kkek.pdf', 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    # Create a Word document
    doc = docx.Document()

    # Loop through each page in the PDF file
    for page_num in range(len(pdf_reader.pages)):
        # Extract the text from the page
        page = pdf_reader.pages[page_num]
        text = page.extract_text()


        # Add the text to the Word document
        if page_num == 0:
            # Add the first page to the document
            doc.add_paragraph(text)
        else:
            # Add subsequent pages as new sections
            doc.add_paragraph()
            paragraph = doc.add_paragraph(text)
    # Save the Word document
    doc.save('example.docx')

    # Close the PDF file
    pdf_file.close()







from django.contrib.auth.models import User,Group
def CreateUser(Cleaned_data):
        if CheckUserExistence(Username = Cleaned_data["pinno"] ):
            newUser = User(username=Cleaned_data["pinno"],is_superuser=False)
            newUser.set_password(Cleaned_data["password"])
            newUser.save()
            AddtoGroup("students",newUser)
            return True
        return False
        

def CheckUserExistence(Username):
    if(not User.objects.filter(username=Username).exists()):
        return True

def AddtoGroup(groupname,user):
    studentgroup = Group.objects.get(name=groupname)
    studentgroup.user_set.add(user)

    
